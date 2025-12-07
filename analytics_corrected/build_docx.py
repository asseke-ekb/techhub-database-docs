from pathlib import Path

from docx import Document


FILES_ORDER = [
    "README.md",
    "01_BUSINESS_DOMAINS.md",
    "02_ENTITIES.md",
    "03_DICTIONARIES.md",
    "04_BUSINESS_PROCESSES.md",
    "05_DATA_MODEL.md",
]


def add_markdown_like(doc: Document, text: str) -> None:
    """Простая вставка Markdown-подобного текста в DOCX (заголовки, код, параграфы)."""
    code_mode = False
    code_lines: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")

        # Тоггл блока кода
        if line.strip().startswith("```"):
            if code_mode:
                # закрываем блок
                doc.add_paragraph("\n".join(code_lines), style="Intense Quote")
                code_lines.clear()
                code_mode = False
            else:
                code_mode = True
            continue

        if code_mode:
            code_lines.append(line)
            continue

        # Заголовки по количеству #
        stripped = line.lstrip()
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text_part = stripped[level:].strip()
            doc.add_heading(text_part, min(level, 4))
            continue

        # Обычные строки
        doc.add_paragraph(line if line else "")

    if code_lines:
        doc.add_paragraph("\n".join(code_lines), style="Intense Quote")


def build_docx():
    base = Path(__file__).parent
    doc = Document()
    doc.add_heading("TechHub – аналитическая документация (актуально)", 0)

    for idx, fname in enumerate(FILES_ORDER, start=1):
        path = base / fname
        if not path.exists():
            continue
        doc.add_page_break()
        doc.add_heading(f"{idx}. {fname}", 1)
        add_markdown_like(doc, path.read_text(encoding="utf-8"))

    output = base / "TechHub_Analytics.docx"
    doc.save(output)
    print(f"Generated: {output}")


if __name__ == "__main__":
    build_docx()
